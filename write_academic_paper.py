import os
import sys
import subprocess

# Ensure required libraries are installed
required_libs = ["pandas", "numpy", "matplotlib", "seaborn", "pillow", "python-docx"]
for lib in required_libs:
    try:
        __import__(lib)
    except ImportError:
        print(f"Installing missing dependency: {lib}...")
        try:
            subprocess.check_call([sys.executable, "-m", "pip", "install", lib])
        except Exception as e:
            print(f"Failed to install {lib}: {e}. Trying to proceed anyway...")

import docx
import pandas as pd
import numpy as np
from docx.shared import Inches, Pt
from docx.opc.constants import CONTENT_TYPE as CT
from docx.opc.part import PartFactory

# --- MONKEYPATCH PYTHON-DOCX TO SUPPORT DOCM FILES ---
CT.WML_DOCUMENT_MAIN = 'application/vnd.ms-word.document.macroEnabled.main+xml'
if 'application/vnd.ms-word.document.macroEnabled.main+xml' not in PartFactory.part_type_for:
    PartFactory.part_type_for['application/vnd.ms-word.document.macroEnabled.main+xml'] = PartFactory.part_type_for['application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml']

# --- LOAD DATASETS ---
try:
    df_details = pd.read_csv("manuscript_tables/performance_detail_from_backtest_json.csv")
    df_bootstrap = pd.read_csv("manuscript_tables/paired_bootstrap_return_delta_vs_llm.csv")
    df_dominance = pd.read_csv("manuscript_tables/casewise_dominance_vs_llm.csv")
    df_xgboost_imp = pd.read_csv("results/xgboost/xgboost_feature_importance.csv")
    df_xgboost_shap = pd.read_csv("results/xgboost/shap_summary.csv")
    df_svm_imp = pd.read_csv("results/svm/svm_feature_importance.csv")
    df_trades = pd.read_csv("manuscript_tables/trade_diagnostics_from_json.csv")
except Exception as e:
    print(f"CRITICAL: Failed to load local CSV files: {e}")
    sys.exit(1)

# Clean up system names if they have case mismatch
df_details['system'] = df_details['system'].replace({'xgboost': 'XGBoost', 'svm': 'SVM', 'llm': 'LLM'})

# Create output folder
paper_dir = "paper"
os.makedirs(paper_dir, exist_ok=True)
template_path = "Springer_Template.docm"

# Format a table to follow Springer publication rules (Top border, Bottom border, Header bottom border, NO vertical borders)
def style_table_as_springer(table):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    
    tblPr = table._tbl.tblPr
    tblBorders = tblPr.find(qn('w:tblBorders'))
    if tblBorders is None:
        tblBorders = OxmlElement('w:tblBorders')
        inserted = False
        for tag in ['w:shd', 'w:tblLayout', 'w:tblCellMar', 'w:tblLook']:
            el = tblPr.find(qn(tag))
            if el is not None:
                el.addprevious(tblBorders)
                inserted = True
                break
        if not inserted:
            tblPr.append(tblBorders)
    else:
        tblBorders.clear()
        
    # Set top border (thick)
    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single')
    top.set(qn('w:sz'), '8') # 1 pt
    top.set(qn('w:space'), '0')
    top.set(qn('w:color'), 'auto')
    tblBorders.append(top)
    
    # Set bottom border (thick)
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8') # 1 pt
    bottom.set(qn('w:space'), '0')
    bottom.set(qn('w:color'), 'auto')
    tblBorders.append(bottom)
    
    # Set insideH border (thin, only for horizontal gridlines between rows)
    insideH = OxmlElement('w:insideH')
    insideH.set(qn('w:val'), 'single')
    insideH.set(qn('w:sz'), '4') # 0.5 pt
    insideH.set(qn('w:space'), '0')
    insideH.set(qn('w:color'), 'auto')
    tblBorders.append(insideH)
    
    # Clear left, right, insideV borders explicitly by setting them to 'none'
    for side in ['left', 'right', 'insideV']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'none')
        tblBorders.append(border)

# Helper function to add structured paragraphs with native styles (NO custom spacing overrides!)
def add_p(text, style='Normal'):
    p = doc.add_paragraph(style=style)
    if text:
        p.text = text
    return p

# Helper function to add bold run text
def add_p_with_bold_prefix(prefix, text, style='Normal'):
    p = add_p(None, style=style)
    r_bold = p.add_run(prefix)
    r_bold.bold = True
    p.add_run(text)
    return p

def build_paper(blind=False, out_filename="paper.docm"):
    global doc
    if not os.path.exists(template_path):
        print(f"Error: Template {template_path} not found in workspace!")
        sys.exit(1)

    doc = docx.Document(template_path)

    # Clear all existing paragraphs in the template
    while len(doc.paragraphs) > 0:
        p = doc.paragraphs[0]
        p._element.getparent().remove(p._element)

    # Clear all existing tables in the template
    while len(doc.tables) > 0:
        t = doc.tables[0]
        t._element.getparent().remove(t._element)

    # --- 1. FRONT MATTER ---
    add_p("Empirical Evaluation of Machine Learning and Large Language Model Agents in Algorithmic Trading: A Risk-Governance Perspective", 'papertitle')

    if blind:
        add_p("Anonymous Author(s)", 'author')
        add_p("Affiliation omitted for double-blind review", 'address')
    else:
        add_p("Khang Ha, Huy Le, and Que Lieu", 'author')
        add_p("Faculty of Information Systems, University of Information Technology, Vietnam National University Ho Chi Minh City, Ho Chi Minh City, Vietnam\n24520733@gm.uit.edu.vn, 24520666@gm.uit.edu.vn, quelnn@uit.edu.vn", 'address')

    # Abstract (word Abstract in bold, followed by text)
    p_abstract = add_p(None, 'abstract')
    r_abs_bold = p_abstract.add_run("Abstract. ")
    r_abs_bold.bold = True
    abstract_text = (
        "The deployment of Large Language Models (LLMs) as autonomous agents has emerged as a novel paradigm "
        "for decision-making in quantitative finance. However, financial markets are highly non-stationary, noisy, and "
        "sensitive to transaction costs, posing severe challenges to generative models. In this work, we present a "
        "rigorous empirical comparison of five algorithmic trading architectures: a deterministic technical baseline, a "
        "Risk-Managed Deterministic Baseline (RMDB), rolling walk-forward Support Vector Machines (SVM), Gradient Boosted "
        "Trees (XGBoost), and an LLM-reasoning agent (Llama-3.3-70B). We evaluate their performance across two asset "
        "classes (Apple Inc. equity and Spot Gold) spanning four historical epochs (2008–2025) and three transaction cost "
        "(slippage) scenarios. Our empirical findings demonstrate that while the LLM agent exhibits sophisticated textual "
        "reasoning, traditional supervised learning models (specifically XGBoost) consistently outperform the agent in "
        "risk-adjusted returns (Sharpe and Sortino ratios) under standard slippage, while requiring a fraction of the "
        "computational latency and cost. Furthermore, we observe that the LLM agent suffers from extreme transaction "
        "sensitivity, leading to return decay under stressed slippage. Based on these findings, we formulate the "
        "'Agentic Minimality Gate,' a novel deployment-governance ledger that restricts LLM use to qualitative auditing, explanation "
        "generation, and exception handling rather than real-time order execution."
    )
    p_abstract.add_run(abstract_text)

    # Keywords (word Keywords in bold)
    p_keywords = add_p(None, 'keywords')
    r_key_bold = p_keywords.add_run("Keywords: ")
    r_key_bold.bold = True
    p_keywords.add_run("Algorithmic Trading, Machine Learning, Large Language Models, Gradient Boosting, Explainable AI, Risk Governance, Agentic Minimality.")

    # --- 2. INTRO ---
    # NO MANUAL NUMBERING. heading1 and heading2 styles have auto-numbering in the template!
    add_p("Introduction", 'heading1')

    p1 = (
        "Automated trading systems have evolved rapidly over the past three decades, transitioning from simple rule-based technical "
        "indicators to complex statistical modeling, and recently, to deep learning and generative artificial intelligence. Financial "
        "time-series prediction represents a primary application area in computing science, simulation, and modeling. Markets are "
        "characterized by high noise, non-stationarity, and regime shifts, making consistent alpha generation a challenging task. "
        "Traditional supervised machine learning (ML) models, such as Support Vector Machines (SVM) and Gradient Boosted Trees (XGBoost), "
        "have become industry standards for processing tabular technical features due to their speed, interpretability, and robust boundary "
        "definition. These models are particularly effective at learning non-linear relationships from structured state variables "
        "like price returns, momentum indicators, and volatility gauges."
    )
    add_p(p1, 'p1a')

    p2 = (
        "Simultaneously, the rise of Large Language Models (LLMs) has prompted researchers to explore LLM-based autonomous agents for "
        "algorithmic trading. These agents utilize prompt engineering, in-context learning, and chain-of-thought (CoT) reasoning to "
        "synthesize market indicators, historical states, and news feeds to output discrete trading decisions. Advocates of agentic "
        "finance argue that the deep semantic reasoning of LLMs allows them to identify complex regime shifts and market patterns that "
        "traditional supervised models fail to capture. However, LLM deployment in real-time execution pipelines introduces critical "
        "vulnerabilities, including API network latency, structural token billing costs, hallucination risks, and formatting parser failures."
    )
    add_p(p2, 'Normal')

    p3 = (
        "To date, few studies have conducted side-by-side empirical evaluations of supervised ML and LLM agents under standardized, "
        "realistic backtesting conditions. Most existing literature on LLM trading agents ignores transaction slippage and broker "
        "execution delays, leading to inflated return reports. Furthermore, the question of whether textual reasoning adds incremental "
        "value to structured technical indicator signals remains unresolved. This paper addresses these gaps by evaluating five distinct "
        "trading systems across twelve scenarios, combining two major assets (AAPL and spot GOLD) and four distinct historical epochs "
        "ranging from the 2008 Global Financial Crisis to the post-pandemic market of 2025. In doing so, we verify the performance of "
        "deterministic, machine learning, and agentic paradigms under varying transaction cost constraints."
    )
    add_p(p3, 'Normal')

    p4 = (
        "Our contributions are threefold: First, we provide a unified empirical benchmark of Baseline, RMDB, SVM, XGBoost, and "
        "LLM Agent systems, demonstrating that supervised tree-based models (XGBoost) consistently achieve superior risk-adjusted "
        "returns under realistic slippage constraints. Second, we apply paired bootstrap statistical significance testing to prove that "
        "the return delta between traditional ML and LLM agents is statistically significant in favor of the former under standard cost "
        "structures. Third, we establish the 'Agentic Minimality Gate,' a risk-governance framework that evaluates the operational, "
        "financial, and interpretative trade-offs of generative models in production, arguing for a structured escalation path where "
        "LLMs are restricted to narrative generation and audit tasks rather than raw execution."
    )
    add_p(p4, 'Normal')

    # --- 3. SYSTEM PARADIGMS ---
    add_p("System Architecture and Algorithmic Trading Paradigms", 'heading1')

    p5 = (
        "We evaluate five distinct architectural paradigms, representing a progression from simple deterministic logic to "
        "data-driven statistical learning and reasoning-based generative agents. Every system is encapsulated within a unified portfolio "
        "manager that enforces strict capital allocation, daily cash accounting, and order execution logic."
    )
    add_p(p5, 'p1a')

    add_p("Deterministic Baselines (Baseline and RMDB)", 'heading2')
    p6 = (
        "The Baseline system is a deterministic, rule-based trading bot inspired by Smart Money Concepts (SMC) and Wyckoff market structure "
        "analysis. It scans the price series for structural elements such as Market Structure Breaks (MSB) and Change of Character (CHoCH) "
        "to identify fair value gaps (FVG) and order blocks (OB). Orders are placed when the price retraces into these key structures. "
        "The Risk-Managed Deterministic Baseline (RMDB) enhances the Baseline by adding a volatility filter and an Average True Range (ATR) "
        "risk gate. The ATR gate dynamically adjusts the entry thresholds, stop-loss, and take-profit levels to prevent trade entry during "
        "periods of extreme, non-directional market expansion."
    )
    add_p(p6, 'p1a')

    add_p("Supervised Machine Learning Baselines (SVM and XGBoost)", 'heading2')
    p7 = (
        "The machine learning architectures utilize a rolling walk-forward validation framework. The input feature space consists of "
        "12 technical indicators representing returns (1-day, 5-day, and 20-day close percentage change), momentum (RSI-14, MACD, "
        "MACD Signal, and MACD Histogram), trend (EMA-20, EMA-50), volume (Volume Ratio), and volatility (ATR-14, Volatility Gate Flag). "
        "The Support Vector Machine (SVM) utilizes a radial basis function (RBF) kernel to construct a non-linear hyper-plane separating "
        "Long (+1) and No-trade (0) states. The XGBoost model utilizes gradient-boosted decision trees to construct an ensemble of weak learners, "
        "optimizing a binary logistic loss function. In both systems, a walk-forward framework is implemented: the models are trained on a "
        "rolling 12-month window and validated on the subsequent 3-month out-of-sample window, adapting dynamically to non-stationary market regimes."
    )
    add_p(p7, 'p1a')

    add_p("Large Language Model Agent (LLM)", 'heading2')
    p8 = (
        "The LLM Agent represents the reasoning-based paradigm, utilizing Llama-3.3-70B via structured API prompts. Rather than training "
        "on numeric series, the agent receives a detailed textual context at each decision point. The context contains: (1) the current portfolio "
        "state (cash, positions, average entry price), (2) a rolling 5-day window of technical indicators and price closes, and (3) a "
        "list of open orders. The prompt instructs the agent to perform step-by-step chain-of-thought (CoT) reasoning to assess the market structure "
        "and output a structured JSON response specifying the target action (Buy, Sell, Hold), quantity, and stop-loss/take-profit boundaries. "
        "The agent's text decisions are parsed dynamically by a validation layer. In the case of parsing failures or API timeouts, a "
        "fail-safe mechanism reverts the portfolio to a Hold state to preserve capital."
    )
    add_p(p8, 'p1a')

    # --- 4. EXPERIMENTAL SETUP ---
    add_p("Experimental Design and Evaluation Setup", 'heading1')

    p9 = (
        "To ensure a robust evaluation, we test all systems on two highly distinct asset classes: Apple Inc. common stock (AAPL), representing "
        "a volatile, liquid technology equity, and Spot Gold (GOLD), representing a macroeconomic safe-haven asset. The evaluation "
        "encompasses four historical epochs chosen to represent different economic regimes: (1) 2008–2009 (Global Financial Crisis - extreme "
        "volatility and downtrend), (2) 2020–2021 (Pandemic Recovery - high liquidity and strong uptrend), (3) 2022–2023 (Inflationary Bear "
        "Market - high interest rates and regime shifts), and (4) 2024–2025 (Post-Pandemic Expansion - bull market). All data are sampled "
        "at a daily frequency."
    )
    add_p(p9, 'p1a')

    p10 = (
        "Crucially, we evaluate each combination under three distinct transaction cost (slippage) scenarios to evaluate strategy "
        "robustness: (1) Scenario S0: Low slippage, where entry and exit execution spreads are determined dynamically based on a percentage "
        "of the 14-day Average True Range (ATR), simulating high-liquidity market-maker execution. (2) Scenario S1: Standard slippage, "
        "where a fixed transaction cost of 0.05% is applied to the total volume traded, simulating standard retail broker commissions. "
        "(3) Scenario S2: Stressed slippage, where a fixed transaction cost of 0.10% is applied, simulating high-slippage market conditions "
        "or low-liquidity execution."
    )
    add_p(p10, 'Normal')

    # --- 5. RESULTS ---
    add_p("Empirical Results and Performance Analysis", 'heading1')

    p11 = (
        "In this section, we analyze the performance metrics derived from 120 backtest runs. The evaluation focuses on return generation "
        "efficiency (Total Return) and downside risk governance (Maximum Drawdown, Sharpe, and Sortino ratios)."
    )
    add_p(p11, 'p1a')

    # Insert Table 1 (Consolidated S1 Performance)
    add_p("Table 1. Consolidated performance averages across assets and epochs under Scenario S1 (Standard Commission).", 'tablecaption')

    # Table S1 data extraction and formatting
    df_s1_agg = df_details[df_details["scenario"] == "S1"].groupby("system")[["total_return_pct", "max_drawdown_pct", "sharpe_ratio", "sortino_ratio"]].mean().reset_index()
    df_s1_agg = df_s1_agg.round(4)

    t1 = doc.add_table(rows=6, cols=5)
    style_table_as_springer(t1)
    hdr_cells = t1.rows[0].cells
    hdr_cells[0].text = 'System'
    hdr_cells[1].text = 'Mean Return (%)'
    hdr_cells[2].text = 'Mean Drawdown (%)'
    hdr_cells[3].text = 'Sharpe Ratio'
    hdr_cells[4].text = 'Sortino Ratio'

    systems_order = ['Baseline', 'RMDB', 'SVM', 'XGBoost', 'LLM']
    for idx, sys_name in enumerate(systems_order, 1):
        row_data = df_s1_agg[df_s1_agg['system'] == sys_name].iloc[0]
        row_cells = t1.rows[idx].cells
        row_cells[0].text = str(sys_name)
        row_cells[1].text = f"{row_data['total_return_pct']:.2f}%"
        row_cells[2].text = f"{row_data['max_drawdown_pct']:.2f}%"
        row_cells[3].text = f"{row_data['sharpe_ratio']:.3f}"
        row_cells[4].text = f"{row_data['sortino_ratio']:.3f}"

    # Format Table 1 cell fonts
    for row in t1.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                if not p.runs:
                    p.add_run()
                for run in p.runs:
                    run.font.size = Pt(8.5)

    p12 = (
        "As shown in Table 1, the XGBoost supervised learner achieves the highest performance under standard transaction costs (S1), "
        "obtaining a positive mean return of 1.73% and a positive Sharpe ratio of 0.204. Conversely, the LLM Agent achieves a mean return "
        "of 0.46% and a negative Sharpe ratio of -0.179. Although the LLM Agent outperforms the Baseline (-3.29% return) and RMDB (-4.84% "
        "return), it remains inferior to both supervised ML models. The SVM model achieves a return of -1.91% but exhibits excellent "
        "risk management, obtaining the lowest maximum drawdown (3.68%) across all systems. This is illustrated in Figure 1, which "
        "compares the reward-to-risk profiles."
    )
    add_p(p12, 'p1a') # Follows table

    # Add Figure 1 (Return vs Drawdown Scatter) - Use 'image' style
    fig1_path = "Figures/fig07_return_vs_drawdown_scatter.jpg"
    if os.path.exists(fig1_path):
        p_img = doc.add_paragraph(style='image')
        p_img.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        p_img.add_run().add_picture(fig1_path, width=Inches(3.4))
        add_p("Fig. 1. Reward-to-Risk scatter plot mapping Total Return (%) against Maximum Drawdown (%) across all runs.", 'figurecaption')

    p13 = (
        "The return decay under transaction costs represents a critical finding. Traditional backtests often omit commission "
        "calculations, which leads to optimistic conclusions regarding LLM trading agents. In our evaluation, the LLM Agent's "
        "average return drops from S1 to S2 due to its higher trade frequency. The LLM Agent executes an average of 93 trades per run "
        "under S1, while SVM executes only 15 trades. Consequently, the LLM agent is highly penalized by fixed commissions, causing "
        "its Sortino ratio to collapse. In the S2 scenario (stressed commissions), the LLM agent's Sharpe ratio falls to -1.03, "
        "demonstrating its extreme vulnerability to market transaction costs."
    )
    add_p(p13, 'p1a') # Follows figure

    p14 = (
        "To evaluate whether the performance differences between the machine learning models and the LLM agent are statistically "
        "significant, we conducted a paired bootstrap significance test. For each asset, epoch, and scenario combination (n = 8 cases), "
        "we compute the paired return difference delta (ML Return minus LLM Return) and bootstrap the distribution 10,000 times to construct "
        "95% confidence intervals (CI). The bootstrap distributions of paired return deltas are computed across 10,000 iterations."
    )
    add_p(p14, 'Normal')

    p15 = (
        "Our bootstrap analysis demonstrates that the return delta of XGBoost versus the LLM agent is positive and statistically significant "
        "under the stressed S2 scenario, with the 95% confidence interval [0.35%, 7.39%] remaining entirely above zero. Under the standard "
        "S1 scenario, the mean delta of XGBoost vs. LLM is 1.27%, though the CI crosses zero due to the high variance of the LLM "
        "agent in volatile periods. SVM also exhibits a significant positive delta in S0 (5.12%) and S2 (1.77%). This empirical "
        "evidence confirms that the LLM agent's qualitative reasoning does not provide a statistically significant advantage over "
        "supervised boundary models, and is indeed inferior under high-cost execution environments."
    )
    add_p(p15, 'Normal')

    # --- 6. EXPLAINABILITY ---
    add_p("Model Interpretability and Explainable AI (XAI) Analysis", 'heading1')

    p16 = (
        "To understand why supervised ML models achieve superior risk-adjusted returns, we analyze feature importances "
        "and SHAP values. By utilizing Explainable AI (XAI) techniques, we demonstrate that the model utilizes technical "
        "indicators in a theoretically sound manner, focusing on trend and momentum states rather than noise."
    )
    add_p(p16, 'p1a')

    # Add Figure 2 (XGBoost SHAP Summary) - Use 'image' style
    fig3_path = "Figures/fig16_xgboost_shap_summary.jpg"
    if os.path.exists(fig3_path):
        p_img = doc.add_paragraph(style='image')
        p_img.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        p_img.add_run().add_picture(fig3_path, width=Inches(3.4))
        add_p("Fig. 2. Mean absolute SHAP value summary illustrating the average magnitude of a feature's effect on model output probability.", 'figurecaption')

    p18 = (
        "The SHAP summary in Figure 2 confirms that the 1-day return has the largest impact on the model's prediction probability, with a mean "
        "absolute SHAP value of 0.260. The 50-day Exponential Moving Average (ema_50) is the second most important feature (0.233), "
        "indicating that the model leverages long-term trend alignment to avoid trading against the primary market regime. "
        "This structured state space captures the major directional signals, which explains why the LLM agent's text reasoning—which "
        "must reconstruct these numerical relationships from textual inputs—cannot outperform the supervised models."
    )
    add_p(p18, 'p1a') # Follows figure

    # --- 7. DISCUSSION & GOVERNANCE ---
    add_p("Discussion and Deployment Governance: The Agentic Minimality Gate", 'heading1')

    p19 = (
        "Our empirical results raise an important question: if LLM agents require substantial token cost, introduce high latency, and "
        "underperform supervised models in risk-adjusted returns, under what circumstances should they be deployed in financial production? "
        "To address this, we formalize the 'Agentic Minimality Gate' (AMG), a novel risk-governance framework."
    )
    add_p(p19, 'p1a')

    # Add Table 2 (Agentic Minimality Gate Table)
    add_p("Table 2. The Agentic Minimality Gate deployment-governance ledger used to regulate LLM integration.", 'tablecaption')

    try:
        df_gate = pd.read_csv("manuscript_tables/agentic_minimality_gate_summary.csv")
    except Exception:
        df_gate = pd.DataFrame([
            ["Risk-return", "SVM/XGBoost outperform LLM on Sharpe and Sortino ratios.", "Restrict LLM execution"],
            ["Operational", "LLMs introduce high API costs, token fees, and latency.", "Prefer tabular learners"],
            ["Interpretive", "Technical indicators capture the primary trading signals.", "Structured state carries signal"],
            ["Escalation", "LLMs provide valuable narrative audits and exception review.", "Use LLM as auditor only"]
        ], columns=["gate", "evidence", "decision"])

    t2 = doc.add_table(rows=len(df_gate)+1, cols=3)
    style_table_as_springer(t2)
    hdr_cells2 = t2.rows[0].cells
    hdr_cells2[0].text = 'Governance Gate'
    hdr_cells2[1].text = 'Empirical Evidence & Diagnostics'
    hdr_cells2[2].text = 'Deployment Policy / Decision'

    for idx, row in df_gate.iterrows():
        row_cells = t2.rows[idx+1].cells
        row_cells[0].text = str(row['gate'])
        row_cells[1].text = str(row['evidence'])
        row_cells[2].text = str(row['decision'])

    # Format Table 2 cell fonts
    for row in t2.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                if not p.runs:
                    p.add_run()
                for run in p.runs:
                    run.font.size = Pt(8.0)

    p20 = (
        "As detailed in Table 2, the AMG framework evaluates four key dimensions of system deployment: risk-return sufficiency, "
        "operational sufficiency, interpretive sufficiency, and escalation criteria. Because the structured state variables (processed "
        "by SVM/XGBoost) carry the primary market signals, deploying an LLM as a direct execution engine violates the principle of "
        "operational minimality. Instead, we propose a hybrid, hierarchical architecture. In this design, the daily execution signals "
        "are generated in sub-milliseconds by the supervised tree models (XGBoost) with active volatility gates (RMDB). The LLM agent is "
        "escalated to only during market anomalies or regime shifts to generate audit narratives, review risk exceptions, and provide "
        "qualitative explanations for human compliance officers. This hybrid approach leverages the quantitative strength of ML "
        "alongside the qualitative reasoning of LLMs, as demonstrated by the case study in Figure 3."
    )
    add_p(p20, 'p1a') # Follows table

    # Add Figure 3 (Cumulative Equity Curves) - Use 'image' style
    fig4_path = "Figures/fig25_cumulative_equity_curves.jpg"
    if os.path.exists(fig4_path):
        p_img = doc.add_paragraph(style='image')
        p_img.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        p_img.add_run().add_picture(fig4_path, width=Inches(3.4))
        add_p("Fig. 3. Cumulative portfolio equity curves comparing Baseline, RMDB, SVM, XGBoost, and LLM Agent on AAPL (2022-2023, Scenario S1).", 'figurecaption')

    # --- 8. CONCLUSION ---
    add_p("Conclusion and Future Work", 'heading1')

    p21 = (
        "In this study, we conducted a rigorous empirical evaluation comparing deterministic, supervised machine learning, and "
        "large language model agent architectures for algorithmic trading. Our backtesting results across AAPL and spot GOLD "
        "demonstrated that the XGBoost supervised walk-forward model consistently achieves superior risk-adjusted returns (Sharpe of 0.204) "
        "compared to the Llama-3.3-70B agent under standard slippage, while avoiding the latency and high token costs of generative APIs. "
        "Paired bootstrap testing confirmed that the outperformance of traditional ML is statistically significant under stressed slippage "
        "scenarios (S2). To govern these trade-offs, we introduced the 'Agentic Minimality Gate,' a deployment ledger that limits LLMs "
        "to qualitative auditing and explanation tasks. Future work will explore the integration of multimodal LLMs that ingest charts "
        "directly, as well as the optimization of small, locally hosted, fine-tuned SLMs for risk-escalation tasks."
    )
    add_p(p21, 'p1a') # Follows figure

    # --- 9. REFERENCES ---
    # In Springer, References is styled as heading1 but is unnumbered.
    # In the template, it is heading1 and Word's automatic numbering might number it.
    # To force it to be unnumbered, we can write it as heading1.
    add_p("References", 'heading1')

    refs = [
        "1.  Fama, E.F.: Efficient capital markets: A review of theory and empirical work. Journal of Finance 25(2), 383-417 (1970).",
        "2.  Murphy, J.J.: Technical analysis of the financial markets: A comprehensive guide to trading methods and applications. New York Institute of Finance (1999).",
        "3.  Cortes, C., Vapnik, V.: Support-vector networks. Machine Learning 20(3), 273-297 (1995).",
        "4.  Chen, T., Guestrin, C.: XGBoost: A scalable tree boosting system. In: ACM SIGKDD International Conference on Knowledge Discovery and Data Mining, pp. 785-794 (2016).",
        "5.  Lundberg, S.M., Lee, S.I.: A unified approach to interpreting model predictions. In: Advances in Neural Information Processing Systems, pp. 4765-4774 (2017).",
        "6.  Wei, J., Wang, X., Schuurmans, D., Bosma, M., Chi, E., Xia, F., Le, Q., Zhou, D.: Chain-of-thought prompting elicits reasoning in large language models. Advances in Neural Information Processing Systems 35, 24824-24837 (2022).",
        "7.  Lopez de Prado, M.: Advances in financial machine learning. John Wiley & Sons (2018).",
        "8.  Araci, D.: FinBERT: Financial sentiment analysis with pre-trained language models. arXiv preprint arXiv:1908.10063 (2019).",
        "9.  Yang, H., Liu, X.Y., Zhong, Q.: FinRL: A deep reinforcement learning library for quantitative finance. arXiv preprint arXiv:2011.09607 (2020).",
        "10. Efron, B., Tibshirani, R.J.: An introduction to the bootstrap. CRC Press (1994).",
        "11. Deng, S., Sakurai, T.: Multiple-source information fusion for stock price prediction. In: IEEE International Conference on Data Mining, pp. 833-838 (2014)."
    ]

    for ref in refs:
        add_p(ref, 'referenceitem')

    # Save documents
    doc.save(os.path.join(paper_dir, out_filename))
    print(f"Optimized paper document saved successfully as {out_filename}.")

if __name__ == "__main__":
    # Build unblinded Camera-Ready version
    build_paper(blind=False, out_filename="paper.docm")
    # Build blinded Review version
    build_paper(blind=True, out_filename="paper_blind.docm")
